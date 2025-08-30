import logging
import os
import re
from typing import Tuple, Optional

import pdfplumber
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
import openpyxl
from openpyxl.utils.exceptions import InvalidFileException

logger = logging.getLogger(__name__)


class FileParser:
    """文件解析工具类，支持多种格式文件的文本提取（纯Python实现）"""

    def __init__(self):
        """初始化文件解析器"""
        logger.info("初始化文件解析器（增强版）")
        # 定义有意义文本的判断模式
        self.meaningful_text_pattern = re.compile(r'[\u4e00-\u9fa5a-zA-Z0-9]{2,}')

    def parse_file(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """
        解析文件并提取文本内容

        Args:
            file_path: 文件路径

        Returns:
            三元组 (是否成功, 文本内容, 错误信息)
        """
        if not os.path.exists(file_path):
            return False, "", f"文件不存在: {file_path}"

        if not os.path.isfile(file_path):
            return False, "", f"不是有效文件: {file_path}"

        # 获取文件扩展名
        file_ext = os.path.splitext(file_path)[1].lower()
        logger.info(f"解析文件: {file_path}，识别到扩展名: {file_ext}")

        try:
            if file_ext in ['.txt', '.text']:
                success, content, msg = self._parse_text(file_path)
            elif file_ext in ['.pdf']:
                success, content, msg = self._parse_pdf(file_path)
            elif file_ext in ['.docx']:
                success, content, msg = self._parse_docx(file_path)
            elif file_ext in ['.doc']:
                return False, "", ".doc格式暂不支持，请转换为.docx后重试"
            elif file_ext in ['.xlsx', '.xls']:
                success, content, msg = self._parse_excel(file_path)
            else:
                # 尝试文本解析作为 fallback
                success, content, msg = self._parse_text_fallback(file_path)

            # 验证提取的文本是否有意义
            if success and not self._is_meaningful_text(content):
                return False, content, "提取的文本内容可能无实际意义，请检查文件是否为可识别的文本格式"

            return success, content, msg

        except Exception as e:
            error_msg = f"解析文件失败: {str(e)}"
            logger.error(f"解析文件 {file_path} 时发生异常: {error_msg}", exc_info=True)
            return False, "", error_msg

    def _parse_text(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """解析文本文件，支持多种编码尝试"""
        encodings = ['utf-8', 'gbk', 'gb2312', 'iso-8859-1', 'utf-16']

        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as f:
                    content = f.read()
                content = self._clean_text(content)
                return True, content, None
            except UnicodeDecodeError:
                continue
            except Exception as e:
                return False, "", f"读取文本文件失败: {str(e)}"

        return False, "", "无法识别文本文件编码格式"

    def _parse_text_fallback(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """当无法识别文件格式时，尝试作为文本文件解析（fallback机制）"""
        try:
            # 先尝试二进制读取判断是否为文本文件
            with open(file_path, 'rb') as f:
                sample = f.read(1024)
                # 简单判断是否为文本文件（检查不可打印字符比例）
                printable_ratio = len([c for c in sample if 32 <= c <= 126 or c in [9, 10, 13]]) / len(sample)
                if printable_ratio < 0.7:
                    return False, "", f"不支持的文件格式: {os.path.splitext(file_path)[1]}"

            # 尝试解析为文本
            return self._parse_text(file_path)
        except Exception as e:
            return False, "", f"尝试文本解析失败: {str(e)}"

    def _parse_pdf(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """解析PDF文件，优化布局分析"""
        try:
            content = []
            with pdfplumber.open(file_path) as pdf:
                logger.info(f"开始解析PDF文件，共 {len(pdf.pages)} 页")

                for page_num, page in enumerate(pdf.pages, 1):
                    # 优化布局分析参数，提高文本提取质量
                    text = page.extract_text(
                        x_tolerance=2,  # 横向 tolerance，处理文字轻微错位
                        y_tolerance=2,  # 纵向 tolerance
                        layout=True  # 保留布局信息
                    )

                    if text:
                        content.append(f"=== 第 {page_num} 页 ===")
                        content.append(text)
                        logger.debug(f"PDF第 {page_num} 页提取文本长度: {len(text)}")
                    else:
                        logger.warning(f"PDF第 {page_num} 页未提取到文本内容，尝试其他方法")
                        # 尝试提取页面中的字符
                        chars = page.chars
                        if chars:
                            text = ''.join([c['text'] for c in sorted(chars, key=lambda x: (x['y0'], x['x0']))])
                            content.append(f"=== 第 {page_num} 页 (字符模式) ===")
                            content.append(text)

            if not content:
                return False, "", "PDF文件中未提取到任何文本内容"

            full_text = '\n'.join(content)
            full_text = self._clean_text(full_text)

            # 记录提取的文本样本用于调试
            sample_text = full_text[:200] + "..." if len(full_text) > 200 else full_text
            logger.info(f"PDF解析完成，文本样本: {sample_text}")

            return True, full_text, None
        except Exception as e:
            return False, "", f"PDF解析失败: {str(e)}"

    def _parse_docx(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """解析docx文件"""
        try:
            doc = Document(file_path)
            content = []

            # 提取段落文本
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    content.append(text)

            # 提取表格内容
            for table in doc.tables:
                content.append("\n=== 表格开始 ===")
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        content.append('\t'.join(row_text))
                content.append("=== 表格结束 ===\n")

            if not content:
                return False, "", "DOCX文件中未提取到任何文本内容"

            full_text = '\n'.join(content)
            full_text = self._clean_text(full_text)
            return True, full_text, None
        except PackageNotFoundError:
            return False, "", "无效的DOCX文件或文件已损坏"
        except Exception as e:
            return False, "", f"DOCX解析失败: {str(e)}"

    def _parse_excel(self, file_path: str) -> Tuple[bool, str, Optional[str]]:
        """解析Excel文件（.xlsx和.xls格式）"""
        try:
            # 尝试用openpyxl打开
            wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)
            content = []

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]
                content.append(f"\n=== 工作表: {sheet_name} ===")

                # 读取前100行数据（防止过大文件）
                row_count = 0
                for row in sheet.iter_rows(values_only=True):
                    row_count += 1
                    if row_count > 100:
                        content.append("... 已省略部分内容 ...")
                        break

                    row_text = [str(cell) if cell is not None else "" for cell in row]
                    row_text = [t.strip() for t in row_text if t.strip()]
                    if row_text:
                        content.append('\t'.join(row_text))

            wb.close()

            if not content:
                return False, "", "Excel文件中未提取到任何数据"

            full_text = '\n'.join(content)
            full_text = self._clean_text(full_text)
            return True, full_text, None
        except InvalidFileException:
            return False, "", "无效的Excel文件或文件已损坏"
        except Exception as e:
            return False, "", f"Excel解析失败: {str(e)}"

    def _clean_text(self, text: str) -> str:
        """优化文本清理逻辑，避免过度过滤有意义的内容"""
        # 去除多余的空白字符，但保留段落结构
        text = re.sub(r'\s+', ' ', text)

        # 只移除连续的特殊字符，保留单个特殊字符
        text = re.sub(r'([^\w\s]){2,}', r'\1', text)

        # 修复可能的断词（如"Go-ogle"改为"Google"）
        text = re.sub(r'(\w+)-\s+(\w+)', r'\1\2', text)

        # 保留英文缩写中的点号（如U.S.A.）
        text = re.sub(r'(?<=[A-Z])\.(?=[A-Z])', '.', text)

        # 去除首尾空白
        return text.strip()

    def _is_meaningful_text(self, text: str) -> bool:
        """判断文本是否包含有意义的内容"""
        # 基本长度检查
        if len(text) < 100:
            return False

        # 检查是否有足够的有意义字符序列
        meaningful_matches = self.meaningful_text_pattern.findall(text)
        if len(meaningful_matches) < 10:
            return False

        # 检查标点符号比例，避免全是标点的情况
        punctuation_ratio = len(re.findall(r'[^\w\s]', text)) / max(len(text), 1)
        if punctuation_ratio > 0.3:  # 标点符号占比不超过30%
            return False

        return True
