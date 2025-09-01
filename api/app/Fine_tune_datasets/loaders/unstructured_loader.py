import os

import docx
from unstructured.partition.auto import partition

from app.Fine_tune_datasets.foundation.data_processor import DataLoader


class UnstructuredDataLoader(DataLoader):
    """非结构化数据加载器，支持PDF、Word、TXT等格式"""

    def load(self, source_info):
        """
        加载非结构化数据

        Args:
            source_info: 包含文件路径和格式的字典

        Returns:
            数据列表，每页或每段作为一个数据项
        """
        file_path = source_info.get('path')
        file_format = source_info.get('format', self._infer_format(file_path))

        if not file_path or not os.path.exists(file_path):
            raise ValueError(f"文件路径不存在: {file_path}")

        load_methods = {
            'pdf': self._load_pdf,
            'docx': self._load_docx,
            'doc': self._load_docx,
            'txt': self._load_text,
            'auto': self._load_auto  # 自动检测格式
        }

        if file_format not in load_methods:
            raise ValueError(f"不支持的文件格式: {file_format}")

        return load_methods[file_format](file_path)

    def _infer_format(self, file_path):
        """从文件路径推断文件格式"""
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')
        return ext if ext else 'auto'

    def _load_pdf(self, file_path):
        """加载PDF文件（使用unstructured库提取）"""
        try:
            from turtledemo.sorting_animate import partition
            elements = partition(filename=file_path)
            pages = []
            for i, elem in enumerate(elements):
                text = str(elem).strip()  # 确保内容为字符串并去空
                if text:
                    pages.append({
                        'page_number': i + 1,  # 按元素顺序标记页码（可根据实际需求调整）
                        'content': text,
                        'file_path': file_path
                    })
            return pages
        except Exception as e:
            raise ValueError(f"PDF解析失败: {str(e)}")

    def _load_docx(self, file_path):
        """加载Word文档"""
        try:
            doc = docx.Document(file_path)
            full_text = []

            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)

            content = "\n".join(full_text)

            # 按段落分割
            paragraphs = []
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():
                    paragraphs.append({
                        'paragraph_number': i + 1,
                        'content': para.text,
                        'full_content': content,
                        'file_path': file_path
                    })

            return paragraphs

        except Exception as e:
            raise ValueError(f"Word文档解析失败: {str(e)}")

    def _load_text(self, file_path):
        """加载文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()

            # 按段落分割
            paragraphs = []
            lines = content.split('\n')
            for i, line in enumerate(lines):
                if line.strip():
                    paragraphs.append({
                        'line_number': i + 1,
                        'content': line.strip(),
                        'full_content': content,
                        'file_path': file_path
                    })

            return paragraphs

        except Exception as e:
            raise ValueError(f"文本文件解析失败: {str(e)}")

    def _load_auto(self, file_path):
        """自动检测文件格式并加载"""
        ext = os.path.splitext(file_path)[1].lower().lstrip('.')

        if ext in ['pdf']:
            return self._load_pdf(file_path)
        elif ext in ['docx', 'doc']:
            return self._load_docx(file_path)
        elif ext in ['txt', 'md']:
            return self._load_text(file_path)
        else:
            raise ValueError(f"不支持自动解析的文件格式: {ext}")